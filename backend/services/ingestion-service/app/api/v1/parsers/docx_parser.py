import io
import re
from datetime import datetime
from decimal import Decimal
import pandas as pd
from docx import Document

from .base import BaseStatementParser
from .pdf_parser import _detect_currency, _extract_counterparty, _detect_channel, _extract_upi

class DOCXParser(BaseStatementParser):
    def parse(self, file_bytes: bytes) -> tuple[list[dict], list[dict]]:
        valid = []
        invalid = []

        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Find owner account from early paragraphs
            owner_account = "UNKNOWN"
            ac_patterns = [
                r"A[/]?C\s*(?:No|Number|Num)[^:]*:\s*([A-Za-z0-9-]{6,35})",
                r"Account\s*(?:No|Number|Num)[^:]*:\s*([A-Za-z0-9-]{6,35})",
                r"Savings\s+A[/]?C[^:]*:\s*([A-Za-z0-9-]{6,35})",
                r"Current\s+A[/]?C[^:]*:\s*([A-Za-z0-9-]{6,35})",
                r"Account\s+Number[^:]*:\s*([A-Za-z0-9Xx*-]{6,35})",
            ]
            
            all_text = ""
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                all_text += text + "\n"
                if owner_account == "UNKNOWN":
                    for pat in ac_patterns:
                        m = re.search(pat, text, re.IGNORECASE)
                        if m:
                            owner_account = m.group(1).strip()
                            break

            # Parse tables
            for table in doc.tables:
                if len(table.rows) < 2:
                    continue
                
                header = [str(cell.text or "").lower() for cell in table.rows[0].cells]
                
                date_idx    = next((i for i, h in enumerate(header) if "date" in h or "value dt" in h), -1)
                narr_idx    = next((i for i, h in enumerate(header) if "narration" in h or "description" in h or "particulars" in h or "remark" in h or "purpose" in h), -1)
                debit_idx   = next((i for i, h in enumerate(header) if h in ("debit", "dr", "withdrawal", "debit amount", "debit (dr.)")), -1)
                credit_idx  = next((i for i, h in enumerate(header) if h in ("credit", "cr", "deposit", "credit amount", "credit (cr.)")), -1)
                amount_idx  = next((i for i, h in enumerate(header) if "amount" in h and debit_idx == -1), -1)
                type_idx    = next((i for i, h in enumerate(header) if h in ("type", "transaction type", "dr/cr", "cr/dr", "txn type")), -1)
                balance_idx = next((i for i, h in enumerate(header) if "balance" in h), -1)
                ref_idx     = next((i for i, h in enumerate(header) if "ref" in h or "chq" in h or "cheque" in h or "txn id" in h), -1)
                sender_idx   = next((i for i, h in enumerate(header) if "sender" == h), -1)
                receiver_idx = next((i for i, h in enumerate(header) if "receiver" == h), -1)

                is_indian_format = debit_idx != -1 or credit_idx != -1

                for r_idx, row in enumerate(table.rows[1:]):
                    cells = [c.text for c in row.cells]
                    if not cells or all(c == "" for c in cells):
                        continue
                        
                    try:
                        ts_str = str(cells[date_idx] or "") if date_idx != -1 and date_idx < len(cells) else ""
                        try:
                            timestamp = pd.to_datetime(ts_str, dayfirst=True).to_pydatetime()
                        except Exception:
                            timestamp = datetime.utcnow()

                        debit_val  = 0.0
                        credit_val = 0.0
                        raw_debit = str(cells[debit_idx] or "") if debit_idx != -1 and debit_idx < len(cells) else ""
                        raw_credit = str(cells[credit_idx] or "") if credit_idx != -1 and credit_idx < len(cells) else ""
                        raw_amount = str(cells[amount_idx] or "") if amount_idx != -1 and amount_idx < len(cells) else ""

                        if is_indian_format:
                            if debit_idx != -1 and debit_idx < len(cells):
                                cleaned = raw_debit.replace(",", "").replace("₹", "").strip()
                                debit_val = float(cleaned) if cleaned and cleaned not in ("-", "") else 0.0
                            if credit_idx != -1 and credit_idx < len(cells):
                                cleaned = raw_credit.replace(",", "").replace("₹", "").strip()
                                credit_val = float(cleaned) if cleaned and cleaned not in ("-", "") else 0.0
                        elif amount_idx != -1 and amount_idx < len(cells):
                            cleaned = raw_amount.replace(",", "").replace("₹", "").replace("$", "").strip()
                            try:
                                val = float(cleaned)
                                if type_idx != -1 and type_idx < len(cells):
                                    txn_type_str = str(cells[type_idx] or "").lower()
                                    if "dr" in txn_type_str or "debit" in txn_type_str or "withdrawal" in txn_type_str:
                                        debit_val = abs(val)
                                    else:
                                        credit_val = abs(val)
                                else:
                                    if val < 0:
                                        debit_val = abs(val)
                                    else:
                                        credit_val = val
                            except ValueError:
                                continue

                        amount = debit_val if debit_val > 0 else credit_val
                        if amount <= 0:
                            continue

                        narration = str(cells[narr_idx] or "") if narr_idx != -1 and narr_idx < len(cells) else "Not Found"
                        ref_id = str(cells[ref_idx] or "") if ref_idx != -1 and ref_idx < len(cells) else "Not Found"
                        
                        balance = None
                        bal_raw = str(cells[balance_idx] or "") if balance_idx != -1 and balance_idx < len(cells) else ""
                        if balance_idx != -1 and balance_idx < len(cells):
                            bal_cleaned = bal_raw.replace(",", "").replace("₹", "").replace("$", "").strip()
                            try:
                                balance = Decimal(bal_cleaned)
                            except:
                                pass
                                
                        upi_id = _extract_upi(narration)

                        if debit_val > 0:
                            txn_type = "DEBIT"
                        else:
                            txn_type = "CREDIT"
                            
                        if sender_idx != -1 and receiver_idx != -1 and sender_idx < len(cells) and receiver_idx < len(cells):
                            sender_account = str(cells[sender_idx] or "") or "Not Found"
                            receiver_account = str(cells[receiver_idx] or "") or "Not Found"
                        else:
                            counterparty = _extract_counterparty(narration)
                            if counterparty == "Not Found" and narration != "Not Found":
                                counterparty = narration
                            if txn_type == "DEBIT":
                                sender_account   = owner_account
                                receiver_account = counterparty
                            else:
                                sender_account   = counterparty
                                receiver_account = owner_account

                        valid.append({
                            "sender_account":   sender_account,
                            "receiver_account": receiver_account,
                            "amount":           Decimal(str(amount)),
                            "balance":          balance,
                            "currency":         _detect_currency(" ".join(str(c) for c in cells if c)) or "INR",
                            "timestamp":        timestamp,
                            "transaction_type": txn_type,
                            "payment_channel":  _detect_channel(narration),
                            "ifsc":             "Not Found",
                            "bank_name":        "Not Found",
                            "branch":           "Not Found",
                            "beneficiary":      receiver_account,
                            "purpose":          narration if narration != "Not Found" else "Not Found",
                            "transaction_id":   ref_id,
                            "upi_id":           upi_id,

                            # Raw verbatim fields
                            "sender_account_raw": sender_account,
                            "receiver_account_raw": receiver_account,
                            "amount_raw": raw_debit or raw_credit or raw_amount,
                            "balance_raw": bal_raw,
                            "timestamp_raw": ts_str,
                            "transaction_id_raw": ref_id,
                            "upi_id_raw": upi_id,
                            "narration_raw": narration,
                        })
                    except Exception as e:
                        invalid.append({"row": r_idx + 2, "data": cells, "reason": str(e)})

        except Exception as e:
            raise ValueError(f"DOCX Parse Exception: {str(e)}")

        return valid, invalid
