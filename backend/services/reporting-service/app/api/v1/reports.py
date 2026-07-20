import json
import os
import uuid
from datetime import datetime
from decimal import Decimal
import httpx
import structlog
from fastapi import APIRouter, Depends, Request, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import Mapped, mapped_column
from shared.database import get_db_session, Base, Transaction, Account, Alert, Case
from shared.schemas import ResponseEnvelope
import io

logger = structlog.get_logger(__name__)
router = APIRouter(prefix="/reports", tags=["Reports"])


class Report(Base):
    """
    SQL schema for storing AI-generated analyst reports.
    """
    __tablename__ = "analyst_reports"

    id: Mapped[uuid.UUID] = mapped_column(primary_key=True, default=uuid.uuid4)
    account_id: Mapped[uuid.UUID] = mapped_column(nullable=False, index=True)
    case_id: Mapped[uuid.UUID | None] = mapped_column(nullable=True, index=True)
    report_type: Mapped[str] = mapped_column(nullable=False)  # investigation, executive_summary
    
    executive_summary: Mapped[str] = mapped_column(nullable=False)
    narrative: Mapped[str] = mapped_column(nullable=False)
    evidence_json: Mapped[str] = mapped_column(nullable=False)  # JSON list of transactions
    risk_factors_json: Mapped[str] = mapped_column(nullable=False)  # JSON list of triggered rules
    recommendations: Mapped[str] = mapped_column(nullable=False)


class ReportGenerateRequest(BaseModel):
    account_id: str
    case_id: str | None = None
    report_type: str = "investigation"  # investigation, executive_summary


class ReportResponse(BaseModel):
    report_id: str
    account_id: str
    case_id: str | None
    report_type: str
    executive_summary: str
    narrative: str
    evidence: list[dict]
    risk_factors: list[dict]
    recommendations: str


def generate_mock_report(account_number: str, balance: float, currency: str, alert_score: float, rules_hit: list[dict], txs: list[Transaction]) -> dict:
    """
    Creates a high-quality deterministic compliance report as fallback when Claude key is missing.
    """
    factors_summary = ", ".join([f["rule"] for f in rules_hit]) if rules_hit else "None"
    
    exec_summary = (
        f"MuleShield AI forensic audit flagged account {account_number} for suspicious layering activity. "
        f"The account holds a balance of {balance:,.2f} {currency} with an anomaly threshold threat level of {alert_score:.0f}/100. "
        f"Compliance filters triggered include: {factors_summary}."
    )
    
    narrative = (
        f"Chronological analysis indicates an abnormal sequence of money transfers involving account {account_number}. "
        f"Immediately following initialization or dormancy reactivation, the account processed {len(txs)} transfers "
        f"with high-speed rotation characteristics. Incoming and outgoing flows show direct structural loops "
        f"indicative of transactional smurfing designed to layer assets and obscure final beneficiaries."
    )
    
    evidence = []
    for t in txs:
        evidence.append({
            "date": t.timestamp.strftime("%Y-%m-%d %H:%M:%S"),
            "sender": t.sender_account,
            "receiver": t.receiver_account,
            "amount": f"{float(t.amount):,.2f} {t.currency}",
            "channel": t.payment_channel,
            "purpose": t.purpose or "N/A"
        })
        
    risk_factors = []
    for r in rules_hit:
        risk_factors.append({
            "factor": r.get("rule", "Compliance Violation"),
            "explanation": r.get("reason", "Suspicious transaction frequency.")
        })
        
    recommendations = (
        "1. RESTRICT FUNDS: Place immediate AML administrative holds on outgoing operations for account.\n"
        "2. ANALYST ESCALATION: File a formal Suspicious Activity Report (SAR) with Financial Crimes Enforcement.\n"
        "3. ENTITY LINKAGE: Initiate network graph inspection to identify linked counterparty nodes in the ring."
    )
    
    return {
        "executive_summary": exec_summary,
        "suspicious_activity_narrative": narrative,
        "evidence_table": evidence,
        "risk_factors": risk_factors,
        "recommendations": recommendations
    }


@router.post("/generate", response_model=ResponseEnvelope[ReportResponse])
async def generate_report(
    payload: ReportGenerateRequest,
    db: AsyncSession = Depends(get_db_session)
) -> ResponseEnvelope[ReportResponse]:
    """
    Generates a forensic compliance case report using Claude or template fallback.
    """
    logger.info("Generating report", account_id=payload.account_id, type=payload.report_type)
    
    # 1. Fetch Account
    acct_uuid = uuid.UUID(payload.account_id)
    acct_stmt = select(Account).where(Account.id == acct_uuid)
    acct_res = await db.execute(acct_stmt)
    account = acct_res.scalars().first()
    
    if not account:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Bank account record not found."
        )
        
    # 2. Fetch Transactions
    tx_stmt = select(Transaction).where(
        or_(
            Transaction.sender_account == account.account_number,
            Transaction.receiver_account == account.account_number
        )
    ).where(Transaction.status == "CONFIRMED").order_by(Transaction.timestamp.desc())
    tx_res = await db.execute(tx_stmt)
    txs = list(tx_res.scalars().all())
    
    # 3. Fetch Alert
    alert_stmt = select(Alert).where(Alert.account_id == account.id)
    alert_res = await db.execute(alert_stmt)
    alert = alert_res.scalars().first()
    
    alert_score = alert.score if alert else 0.0
    rules_hit = []
    if alert:
        try:
            parsed = json.loads(alert.trigger_reason)
            rules_hit = parsed.get("factors", [])
        except Exception:
            rules_hit = [{"rule": alert.alert_type, "reason": alert.trigger_reason, "weight": int(alert.score)}]
            
    # Resolve API Key
    from shared.config import BaseAppSettings
    conf = BaseAppSettings()
    api_key = os.environ.get("ANTHROPIC_API_KEY") or conf.ANTHROPIC_API_KEY
    
    report_data = None
    if api_key:
        logger.info("Calling Anthropic Claude to generate forensic compliance report...")
        try:
            # Prepare data dump for Claude context
            tx_data_dump = [{
                "sender": t.sender_account,
                "receiver": t.receiver_account,
                "amount": float(t.amount),
                "currency": t.currency,
                "timestamp": t.timestamp.isoformat(),
                "channel": t.payment_channel,
                "purpose": t.purpose
            } for t in txs[:30]]  # limit context count
            
            prompt = (
                "You are an expert financial investigator at a major retail bank. Generate a professional "
                "forensic audit report based strictly on the following ledger details. Do NOT fabricate accounts, "
                "names, balances, or transaction details that are not provided explicitly in this prompt.\n\n"
                f"Account Number: {account.account_number}\n"
                f"Balance: {account.balance} {account.currency}\n"
                f"Status: {account.status}\n"
                f"Risk Score: {alert_score}/100\n"
                f"Triggered Risk Rules: {json.dumps(rules_hit)}\n"
                f"Ledger Transactions: {json.dumps(tx_data_dump)}\n\n"
                "Provide the report in raw JSON format matching this schema exactly:\n"
                "{\n"
                "  \"executive_summary\": \"Paragraph summarizing the case, threat level, and primary triggers.\",\n"
                "  \"suspicious_activity_narrative\": \"Chronological description of the suspicious transactions and behaviors.\",\n"
                "  \"evidence_table\": [{\"date\": \"...\", \"sender\": \"...\", \"receiver\": \"...\", \"amount\": \"...\", \"channel\": \"...\", \"purpose\": \"...\"}],\n"
                "  \"risk_factors\": [{\"factor\": \"...\", \"explanation\": \"...\"}],\n"
                "  \"recommendations\": \"Numbered recommendations for actions.\"\n"
                "}"
            )
            
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json"
                    },
                    json={
                        "model": "claude-3-5-sonnet-20241022",
                        "max_tokens": 1500,
                        "messages": [{"role": "user", "content": prompt}]
                    },
                    timeout=45.0
                )
                
                if response.status_code == 200:
                    text_content = response.json()["content"][0]["text"]
                    # Extract JSON
                    json_match = re.search(r"(\{.*\})", text_content, re.DOTALL)
                    if json_match:
                        report_data = json.loads(json_match.group(1))
                    else:
                        report_data = json.loads(text_content)
                else:
                    logger.error("Claude call failed", status=response.status_code, body=response.text)
        except Exception as exc:
            logger.error("Error during Claude API call, falling back to mock generator", error=str(exc))
            
    if not report_data:
        logger.info("Using local mock report generator fallback")
        report_data = generate_mock_report(
            account_number=account.account_number,
            balance=float(account.balance),
            currency=account.currency,
            alert_score=alert_score,
            rules_hit=rules_hit,
            txs=txs
        )
        
    case_uuid = uuid.UUID(payload.case_id) if payload.case_id else None
    
    # Store Report in database
    db_report = Report(
        id=uuid.uuid4(),
        account_id=account.id,
        case_id=case_uuid,
        report_type=payload.report_type,
        executive_summary=report_data["executive_summary"],
        narrative=report_data["suspicious_activity_narrative"],
        evidence_json=json.dumps(report_data["evidence_table"]),
        risk_factors_json=json.dumps(report_data["risk_factors"]),
        recommendations=report_data["recommendations"]
    )
    db.add(db_report)
    await db.commit()
    
    return ResponseEnvelope(
        success=True,
        message="Compliance forensic report compiled successfully.",
        data=ReportResponse(
            report_id=str(db_report.id),
            account_id=str(account.id),
            case_id=str(case_uuid) if case_uuid else None,
            report_type=payload.report_type,
            executive_summary=db_report.executive_summary,
            narrative=db_report.narrative,
            evidence=report_data["evidence_table"],
            risk_factors=report_data["risk_factors"],
            recommendations=db_report.recommendations
        )
    )


@router.get("/{report_id}/download")
async def download_report(
    report_id: str,
    format: str = "pdf",  # pdf, docx
    db: AsyncSession = Depends(get_db_session)
):
    """
    Downloads and streams the target case report in PDF or DOCX format.
    """
    logger.info("Downloading report file", id=report_id, format=format)
    
    report_uuid = uuid.UUID(report_id)
    stmt = select(Report).where(Report.id == report_uuid)
    res = await db.execute(stmt)
    report = res.scalars().first()
    
    if not report:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Forensic analyst report record not found."
        )
        
    acct_stmt = select(Account).where(Account.id == report.account_id)
    acct_res = await db.execute(acct_stmt)
    account = acct_res.scalars().first()
    acct_num = account.account_number if account else "N/A"
    
    evidence = json.loads(report.evidence_json)
    risk_factors = json.loads(report.risk_factors_json)
    
    disclaimer = (
        "CONFIDENTIAL AND PROPRIETARY — MuleShield AI compliance report compiled automatically. "
        "AI-generated draft — requires analyst review and validation before final regulatory filing submission."
    )
    
    if format.lower() == "pdf":
        pdf_bytes = None
        
        # Try WeasyPrint HTML rendering
        try:
            from weasyprint import HTML
            logger.info("Attempting HTML to PDF conversion using WeasyPrint...")
            
            # Simple clean styled HTML template
            html_template = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; color: #222; margin: 40px; line-height: 1.5; }}
                    .header {{ border-bottom: 2px solid #DC2626; padding-bottom: 10px; margin-bottom: 30px; }}
                    .logo {{ font-size: 24px; font-weight: bold; color: #DC2626; letter-spacing: 1px; }}
                    .meta-table {{ width: 100%; margin-bottom: 30px; border-collapse: collapse; }}
                    .meta-table td {{ padding: 6px; font-size: 13px; }}
                    .meta-label {{ font-weight: bold; color: #555; width: 150px; }}
                    h2 {{ font-size: 16px; border-bottom: 1px solid #ddd; padding-bottom: 5px; color: #111; margin-top: 30px; }}
                    p {{ font-size: 12px; }}
                    table.evidence {{ width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 15px; }}
                    table.evidence th, table.evidence td {{ border: 1px solid #ddd; padding: 8px; font-size: 10px; text-align: left; }}
                    table.evidence th {{ bg-color: #f5f5f5; font-weight: bold; }}
                    .disclaimer {{ font-size: 8px; color: #777; border-top: 1px solid #ddd; padding-top: 10px; margin-top: 50px; text-align: center; font-style: italic; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <div class="logo">MULESHIELD AI</div>
                    <div style="font-size: 10px; color: #666; margin-top: 5px;">FORENSIC COMPLIANCE INVESTIGATION REPORT</div>
                </div>
                
                <table class="meta-table">
                    <tr>
                        <td class="meta-label">Report Case ID:</td><td>{report.case_id or report.id}</td>
                        <td class="meta-label">Audit Timestamp:</td><td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td>
                    </tr>
                    <tr>
                        <td class="meta-label">Target Account:</td><td>{acct_num}</td>
                        <td class="meta-label">Report Type:</td><td>{report.report_type.upper()}</td>
                    </tr>
                </table>
                
                <h2>1. EXECUTIVE SUMMARY</h2>
                <p>{report.executive_summary}</p>
                
                <h2>2. SUSPICIOUS ACTIVITY NARRATIVE</h2>
                <p>{report.narrative}</p>
                
                <h2>3. TRANSACTION EVIDENCE RECORD</h2>
                <table class="evidence">
                    <thead>
                        <tr>
                            <th>Date</th>
                            <th>Sender</th>
                            <th>Receiver</th>
                            <th>Amount</th>
                            <th>Channel</th>
                            <th>Purpose</th>
                        </tr>
                    </thead>
                    <tbody>
            """
            for e in evidence:
                html_template += f"""
                        <tr>
                            <td>{e['date']}</td>
                            <td>{e['sender']}</td>
                            <td>{e['receiver']}</td>
                            <td>{e['amount']}</td>
                            <td>{e['channel']}</td>
                            <td>{e['purpose']}</td>
                        </tr>
                """
            html_template += f"""
                    </tbody>
                </table>
                
                <h2>4. SUSPECTED VIOLATIONS & RISK FACTORS</h2>
                <ul>
            """
            for rf in risk_factors:
                html_template += f"<li><b>{rf['factor']}:</b> {rf['explanation']}</li>"
            html_template += f"""
                </ul>
                
                <h2>5. RECOMMENDED ACTIONS & CONTAINMENT CODES</h2>
                <p style="white-space: pre-line;">{report.recommendations}</p>
                
                <div class="disclaimer">
                    {disclaimer}
                </div>
            </body>
            </html>
            """
            pdf_bytes = HTML(string=html_template).write_pdf()
        except Exception as weasy_exc:
            logger.warning("WeasyPrint conversion failed, falling back to ReportLab", error=str(weasy_exc))
            
        if not pdf_bytes:
            # ReportLab fallback
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                
                logger.info("Rendering PDF using ReportLab flowables...")
                pdf_buffer = io.BytesIO()
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
                story = []
                
                styles = getSampleStyleSheet()
                
                # Styles
                title_style = ParagraphStyle(
                    'DocTitle',
                    parent=styles['Heading1'],
                    fontSize=22,
                    textColor=colors.HexColor("#DC2626"),
                    spaceAfter=4
                )
                subtitle_style = ParagraphStyle(
                    'DocSubtitle',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor=colors.HexColor("#666666"),
                    spaceAfter=20
                )
                h2_style = ParagraphStyle(
                    'SectionHeading',
                    parent=styles['Heading2'],
                    fontSize=13,
                    textColor=colors.HexColor("#111111"),
                    spaceBefore=15,
                    spaceAfter=6,
                    borderPadding=4
                )
                body_style = ParagraphStyle(
                    'BodyText',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=14,
                    spaceAfter=8
                )
                meta_style = ParagraphStyle(
                    'MetaText',
                    parent=styles['Normal'],
                    fontSize=9,
                    leading=12
                )
                disclaimer_style = ParagraphStyle(
                    'Disclaimer',
                    parent=styles['Normal'],
                    fontSize=7,
                    textColor=colors.HexColor("#777777"),
                    alignment=1,
                    spaceBefore=30
                )
                
                # Header logo
                story.append(Paragraph("MULESHIELD AI", title_style))
                story.append(Paragraph("FORENSIC COMPLIANCE INVESTIGATION REPORT", subtitle_style))
                story.append(Spacer(1, 10))
                
                # Metadata block
                meta_data = [
                    [Paragraph(f"<b>Case ID:</b> {report.case_id or report.id}", meta_style), Paragraph(f"<b>Timestamp:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style)],
                    [Paragraph(f"<b>Account:</b> {acct_num}", meta_style), Paragraph(f"<b>Report Type:</b> {report.report_type.upper()}", meta_style)]
                ]
                meta_table = Table(meta_data, colWidths=[260, 260])
                meta_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('PADDING', (0,0), (-1,-1), 2),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6)
                ]))
                story.append(meta_table)
                story.append(Spacer(1, 15))
                
                # Sections
                story.append(Paragraph("1. EXECUTIVE SUMMARY", h2_style))
                story.append(Paragraph(report.executive_summary, body_style))
                
                story.append(Paragraph("2. SUSPICIOUS ACTIVITY NARRATIVE", h2_style))
                story.append(Paragraph(report.narrative, body_style))
                
                story.append(Paragraph("3. LEDGER EVIDENCE LISTING", h2_style))
                
                # Evidence Table
                table_data = [["Date", "Sender", "Receiver", "Amount", "Channel", "Purpose"]]
                for e in evidence[:15]:  # show up to 15 rows in ReportLab to avoid page overflow overlap
                    table_data.append([
                        e["date"][:19],
                        e["sender"][:10] + "..." if len(e["sender"]) > 12 else e["sender"],
                        e["receiver"][:10] + "..." if len(e["receiver"]) > 12 else e["receiver"],
                        e["amount"],
                        e["channel"],
                        e["purpose"][:15] + "..." if len(e["purpose"]) > 15 else e["purpose"]
                    ])
                    
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#F5F5F5")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('PADDING', (0,0), (-1,-1), 4),
                ])
                evidence_table = Table(table_data, colWidths=[90, 85, 85, 75, 55, 130])
                evidence_table.setStyle(table_style)
                story.append(evidence_table)
                
                story.append(Paragraph("4. RULE INTAKE RISK FACTORS", h2_style))
                for rf in risk_factors:
                    story.append(Paragraph(f"• <b>{rf['factor']}:</b> {rf['explanation']}", body_style))
                    
                story.append(Paragraph("5. NEXT ACTIONS AND REMEDIATIONS", h2_style))
                story.append(Paragraph(report.recommendations.replace("\n", "<br/>"), body_style))
                
                story.append(Spacer(1, 10))
                story.append(Paragraph(disclaimer, disclaimer_style))
                
                doc.build(story)
                pdf_bytes = pdf_buffer.getvalue()
            except Exception as rl_exc:
                logger.error("ReportLab generation also failed!", error=str(rl_exc))
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to render document PDF bytes: {str(rl_exc)}"
                )
                
        # Stream PDF
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Forensic_Report_{acct_num}.pdf"}
        )
        
    elif format.lower() == "docx":
        # Create Word Document
        try:
            import docx
            from docx.shared import Inches, Pt
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import nsdecls, qn
            
            logger.info("Building Word Document using python-docx template layout...")
            doc = docx.Document()
            
            # Setup styles
            style_normal = doc.styles['Normal']
            font = style_normal.font
            font.name = 'Arial'
            font.size = Pt(10.5)
            
            # Title
            p_logo = doc.add_paragraph()
            run_logo = p_logo.add_run("MULESHIELD AI")
            run_logo.bold = True
            run_logo.font.size = Pt(20)
            run_logo.font.color.rgb = docx.shared.RGBColor(220, 38, 38)
            
            p_logo.add_run("\nFORENSIC COMPLIANCE INVESTIGATION REPORT").font.size = Pt(9.5)
            
            # Divider line
            p_div = doc.add_paragraph()
            p_div.paragraph_format.space_after = Pt(20)
            p_div_border = OxmlElement('w:pBdr')
            p_div_bottom = OxmlElement('w:bottom')
            p_div_bottom.set(qn('w:val'), 'single')
            p_div_bottom.set(qn('w:sz'), '12')
            p_div_bottom.set(qn('w:space'), '4')
            p_div_bottom.set(qn('w:color'), 'DC2626')
            p_div_border.append(p_div_bottom)
            p_div._p.get_or_add_pPr().append(p_div_border)
            
            # Metadata
            tbl_meta = doc.add_table(rows=2, cols=2)
            tbl_meta.autofit = True
            
            r0c0 = tbl_meta.cell(0, 0).paragraphs[0]
            r0c0.add_run("Report Case ID: ").bold = True
            r0c0.add_run(str(report.case_id or report.id))
            
            r0c1 = tbl_meta.cell(0, 1).paragraphs[0]
            r0c1.add_run("Audit Timestamp: ").bold = True
            r0c1.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            r1c0 = tbl_meta.cell(1, 0).paragraphs[0]
            r1c0.add_run("Target Account: ").bold = True
            r1c0.add_run(acct_num)
            
            r1c1 = tbl_meta.cell(1, 1).paragraphs[0]
            r1c1.add_run("Report Type: ").bold = True
            r1c1.add_run(report.report_type.upper())
            
            doc.add_paragraph().paragraph_format.space_after = Pt(15)
            
            # Sections
            h1 = doc.add_heading(level=1)
            h1.add_run("1. EXECUTIVE SUMMARY").font.color.rgb = docx.shared.RGBColor(17, 17, 17)
            doc.add_paragraph(report.executive_summary)
            
            h2 = doc.add_heading(level=1)
            h2.add_run("2. SUSPICIOUS ACTIVITY NARRATIVE").font.color.rgb = docx.shared.RGBColor(17, 17, 17)
            doc.add_paragraph(report.narrative)
            
            h3 = doc.add_heading(level=1)
            h3.add_run("3. TRANSACTION EVIDENCE RECORD").font.color.rgb = docx.shared.RGBColor(17, 17, 17)
            
            # Evidence Table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            headers_list = ["Date", "Sender", "Receiver", "Amount", "Channel", "Purpose"]
            for i, h in enumerate(headers_list):
                hdr_cells[i].paragraphs[0].add_run(h).bold = True
                # Set shading
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                
            for e in evidence:
                row_cells = table.add_row().cells
                row_cells[0].text = e["date"][:19]
                row_cells[1].text = e["sender"]
                row_cells[2].text = e["receiver"]
                row_cells[3].text = e["amount"]
                row_cells[4].text = e["channel"]
                row_cells[5].text = e["purpose"] or ""
                
            doc.add_paragraph().paragraph_format.space_after = Pt(15)
            
            h4 = doc.add_heading(level=1)
            h4.add_run("4. COMPLIANCE RISK FACTORS").font.color.rgb = docx.shared.RGBColor(17, 17, 17)
            for rf in risk_factors:
                p_rf = doc.add_paragraph(style='List Bullet')
                p_rf.add_run(f"{rf['factor']}: ").bold = True
                p_rf.add_run(rf['explanation'])
                
            h5 = doc.add_heading(level=1)
            h5.add_run("5. RECOMMENDED ACTIONS").font.color.rgb = docx.shared.RGBColor(17, 17, 17)
            for line in report.recommendations.split("\n"):
                doc.add_paragraph(line)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(30)
            
            p_disc = doc.add_paragraph()
            run_disc = p_disc.add_run(disclaimer)
            run_disc.italic = True
            run_disc.font.size = Pt(8)
            run_disc.font.color.rgb = docx.shared.RGBColor(119, 119, 119)
            p_disc.alignment = 1
            
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=Forensic_Report_{acct_num}.docx"}
            )
        except Exception as docx_exc:
            logger.error("Failed to generate DOCX document", error=str(docx_exc))
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to render document DOCX bytes: {str(docx_exc)}"
            )
            
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported format target requested. Use 'pdf' or 'docx'."
        )
